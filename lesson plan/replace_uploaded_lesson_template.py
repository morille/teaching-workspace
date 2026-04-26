from __future__ import annotations

import shutil
import sys

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


def apply_run_style(run, font_name: str, font_size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if font_size is not None:
        run.font.size = Pt(font_size)
    if bold is not None:
        run.bold = bold


def set_paragraph_text(paragraph, text: str) -> None:
    paragraph.text = text
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        apply_run_style(run, "宋体", 10.5, None)


def clear_and_set_segments(paragraph, segments: list[tuple[str, str, float | None, bool | None]]) -> None:
    paragraph.text = ""
    for text, font_name, font_size, bold in segments:
        run = paragraph.add_run(text)
        apply_run_style(run, font_name, font_size, bold)


def fill_existing_paragraphs(cell, lines: list[str], default_font: str = "宋体", default_size: float = 10.5, default_bold: bool | None = None) -> None:
    paragraphs = cell.paragraphs
    if not paragraphs:
        p = cell.add_paragraph()
        paragraphs = cell.paragraphs
    if len(lines) > len(paragraphs):
        raise ValueError(f"Not enough paragraphs in cell: need {len(lines)}, have {len(paragraphs)}")
    for i, paragraph in enumerate(paragraphs):
        clear_and_set_segments(paragraph, [(lines[i] if i < len(lines) else "", default_font, default_size, default_bold)])


def normalize_table_fonts(table, default_font: str = "宋体", default_size: float = 10.5) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    apply_run_style(run, default_font, default_size, run.bold)


def main() -> None:
    if len(sys.argv) != 3:
        raise SystemExit("Usage: replace_uploaded_lesson_template.py <template.docx> <output.docx>")

    src, dst = sys.argv[1], sys.argv[2]
    shutil.copyfile(src, dst)
    doc = Document(dst)

    # Update second-section header while preserving its paragraph formatting.
    if len(doc.sections) > 1 and doc.sections[1].header.paragraphs:
        clear_and_set_segments(
            doc.sections[1].header.paragraphs[0],
            [("吉林动画学院 本科课程教案                                                 2025-2026学年第二学期", "微软雅黑", 9, None)],
        )

    info = doc.tables[1]
    for cell, value in [
        (info.cell(0, 1), "中国文化实验动画创作"),
        (info.cell(1, 1), "动画艺术学院"),
        (info.cell(2, 1), "2025/2026第二学期"),
        (info.cell(3, 1), "动画"),
        (info.cell(4, 1), ""),
        (info.cell(5, 1), ""),
    ]:
        fill_existing_paragraphs(cell, [value], default_font="微软雅黑", default_size=10)
    for row in info.rows:
        for c in row.cells:
            for p in c.paragraphs:
                for r in p.runs:
                    apply_run_style(r, "微软雅黑", 10, r.bold)

    main_table = doc.tables[3]
    fill_existing_paragraphs(main_table.cell(0, 1), ["1"], default_font="宋体", default_size=10.5, default_bold=True)
    fill_existing_paragraphs(main_table.cell(0, 5), ["4"], default_font="宋体", default_size=10.5, default_bold=True)
    title_para = main_table.cell(0, 3).paragraphs[0]
    clear_and_set_segments(
        title_para,
        [
            ("第一章 实验动画概述", "宋体", 10.5, True),
            ("\n", "宋体", 10.5, None),
            ("第一课次：实验动画的定义、特征与分类导入", "宋体", 10.5, False),
        ],
    )

    objective_lines = [
        "一、教学目标",
        "知识目标：",
        "1.能够准确说出实验动画的基本定义；",
        "2.能够概括实验动画区别于主流叙事动画的核心特征；",
        "3.能够说出实验动画的三种基本分类维度：按主题、按风格、按材料；",
        "4.能够结合课堂案例，对实验动画作品进行初步分类判断。",
        "能力目标：",
        "1.具备从形式、材料和表达方式切入分析实验动画作品的能力；",
        "2.具备使用基础术语描述实验动画艺术特征的能力；",
        "3.具备围绕“材料语言测试”实践任务进行初步构思的能力。",
        "素养目标：",
        "1.树立勇于创新、敢于突破惯性表达的创作意识；",
        "2.增强对中国动画学派实验传统的认同感与文化自信；",
        "3.培养严谨观察、主动试错和知行合一的学习习惯。",
        "    二、重点难点：",
        "重点：实验动画的定义、特征及三种分类维度的理解。",
        "难点：实验动画分类维度之间的交叉关系，以及材料语言如何参与动画表达。",
        "三、学情分析",
        "学生对“动画”已有较为直观的认识，但对“实验动画”的理解仍偏感性，容易只关注题材和画风，忽略形式、材料与表达方式之间的关系。",
        "课堂需通过案例对比、分类辨析和材料测试导入，帮助学生建立基础观察框架。",
        "",
    ]
    fill_existing_paragraphs(main_table.cell(2, 0), objective_lines, default_font="宋体", default_size=10.5)
    objective_bold_indexes = {0, 14, 17}
    for idx, paragraph in enumerate(main_table.cell(2, 0).paragraphs):
        for run in paragraph.runs:
            apply_run_style(run, "宋体", 10.5, idx in objective_bold_indexes)

    content_lines = [
        "【课前自主学习】",
        "明确目标，发布任务",
        "在平台或学习群发布课前任务，布置学生通过线上资源进行课前预习。",
        "1.思考什么样的作品可以被称为实验动画；",
        "2.观察实验动画与主流动画的差异；",
        "3.查找1-2个风格或材料较特殊的动画案例。",
        "线上预习，总结归纳",
        "通过线上资源，学生初步形成对实验动画的基础认知，并带着问题进入课堂。",
        "1.理解实验动画的基本定义与探索属性；",
        "2.初步辨认实验动画在主题、风格、材料上的分类方式；",
        "3.思考材料如何影响动画的视觉表现和观看感受。",
        "师生讨论，辅导答疑",
        "通过学生展示了解自学情况，根据学生预习内容进行探讨。",
        "1.什么样的作品会被我们直觉地认为是“实验动画”；",
        "2.实验动画与主流动画最明显的差异在哪里；",
        "3.同一作品是否可能同时具有多种分类属性。",
        "",
        "【课中讲析讨论】",
        "一、课程导入与问题激活",
        "1.说明本课程在本学期中的定位、学习目标和最终成果导向；",
        "2.通过案例对比，引导学生观察主流动画与实验动画的差异；",
        "3.围绕“什么是实验动画”展开课堂讨论。",
        "二、教师讲评，评价学习效果",
        "1.总结学生自学期间了解到的知识点；",
        "2.讨论相应问题；",
        "3.补充学生预习阶段未充分理解的知识点。",
        "三、教师精讲，引导合作探究",
        "实验动画的定义、特征与分类框架",
        "（一）创设情景导入",
        "以实验动画案例进行对比分析，观察作品在形式、风格、材料和表达方式上的差异。",
        "",
        "问题分析与解决办法",
        "1.问题分析",
        "什么是实验动画？实验动画与主流动画有何不同？实验动画可以如何分类？",
        "2.解决办法",
        "（1）主要知识点解析",
        "1.1实验动画的定义与特征",
        "一、实验动画的基本定义：",
        "实验动画强调探索性、创新性和个体表达，不以统一的商业叙事模式为唯一目标。",
        "二、实验动画的核心特征：",
        "1.非单一叙事导向",
        "2.强调材料探索",
        "3.注重形式创新",
        "4.鼓励个体表达",
        "三、实验动画与主流动画的差异：",
        "可从叙事目标、视觉逻辑、技术路径、观看体验等方面进行比较。",
        "1.2实验动画的分类框架",
        "一、按主题分类",
        "可关注作品表现的题材方向、文化议题与情感指向。",
        "二、按风格分类",
        "可观察作品在造型、节奏、构图、色彩和视听组织上的整体风格。",
        "三、按材料分类",
        "重点关注材料在动画表达中的作用，如二维逐帧与数字合成技术的结合、综合材料的应用等。",
        "四、按材料分类中的典型形式",
        "1.剪纸动画",
        "2.实物动画",
        "3.偶动画",
        "4.其他综合材料实验形式",
        "1.3中国动画实验传统与课程衔接",
        "一、中国动画学派的实验探索传统",
        "可简要联系水墨、剪纸、折纸等方向，帮助学生建立中国动画实验传统的认知。",
        "二、本课与后续课程的关系",
        "本课重点建立“定义-特征-分类”的基础框架，后续课程再进入中国文化实验动画的探索与实践创作。",
        "三、材料语言导入",
        "材料不仅是技术媒介，也会影响动画的节奏、肌理和情感表达。",
        "四、课终小结",
        "回顾本节课程内容的重点与难点，总结所讲述的主要内容。",
        "课堂测试",
        "选择题",
        "1.下列哪一项不属于实验动画的核心特征？",
        "2.下列哪一项属于按材料分类的观察角度？",
        "3.同一部实验动画作品是否只能归入一种类型？",
        "",
        "    五、作业布置",
        "内容：",
        "完成“材料语言测试”。",
        "要求：",
        "1.利用纸张、布料、树叶等综合材料，或数字笔刷，制作一段约5秒的“材质运动测试”；",
        "2.不要求完整叙事，但需要体现清晰的实验意识；",
        "3.可附简短文字说明，说明所用材料、表现意图和遇到的问题。",
        "",
        "【课后实践应用】",
        "    一、实践作业",
        "实践内容：",
        "利用身边的综合材料或数字笔刷，完成一段约5秒的“材质运动测试”。",
        "二、实践步骤",
        "1.观察材料，确定表现对象。",
        "从纸张、布料、树叶、笔触等材料中选择一种或几种，记录其肌理和运动特性。",
        "2.确定形式，测试运动效果。",
        "思考材料在镜头中的运动方式、节奏变化和视觉重点，完成基础测试。",
        "3.整理结果，准备下节展示。",
        "保留测试片段，并简要梳理材料选择、表现意图和问题反思，便于下节课交流。",
        "三、知识拓展",
        "材料语言并不是简单地“换一种材料做动画”，而是让材料本身的质感、限制与偶然性参与表达。",
        "当创作者真正关注材料特性时，动画的节奏、运动轨迹和观看感受都会发生变化。",
        "这也是实验动画区别于单一技术执行的重要部分。",
    ]
    fill_existing_paragraphs(main_table.cell(4, 0), content_lines, default_font="宋体", default_size=10.5)
    content_heading_indexes = {
        0, 1, 6, 11, 17, 18, 22, 27, 29, 32, 33, 35, 36, 38, 40, 45, 46, 47, 49,
        52, 57, 59, 61, 62, 63, 69, 70, 71, 77, 78, 81, 87
    }
    for idx, paragraph in enumerate(main_table.cell(4, 0).paragraphs):
        is_heading = idx in content_heading_indexes
        size = 12 if is_heading else 10.5
        for run in paragraph.runs:
            apply_run_style(run, "宋体", size, True if is_heading else False)

    design_lines = [
        "",
        "【启发式教学】",
        "教师：课前发布导学任务单，以问题驱动方式引导学生预习“实验动画”的相关内容，激发学生学习兴趣。",
        "",
        "",
        "",
        "学生：完成资料查找、案例收集与预习思考，梳理自己对实验动画的初步理解。",
        "",
        "",
        "教师：在线答疑，了解学生预习情况，为课堂讲授做准备。",
        "",
        "",
        "【课中180分钟】",
        "",
        "【交互式教学】",
        "",
        "【20分钟】课程导入与问题激活。",
        "教师：通过案例对比和问题引导，激活学生已有经验，帮助学生进入“实验动画概述”的学习语境。",
        "",
        "【讲解法】",
        "",
        "【45分钟】实验动画的定义与特征。",
        "",
        "教师：系统讲解实验动画的基本定义与核心特征，并与主流动画进行比较，帮助学生建立清晰概念。",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "【课程思政设计】",
        "",
        "教师：在讲解实验精神时，引导学生认识创新不是脱离传统，而是在理解传统基础上的创造性转化。",
        "",
        "教师：结合动画创作实践，引导学生理解实验探索背后的文化责任与创作者意识。",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "教师：引导学生分析实验动画案例，并从主题、风格、材料三个角度总结分类依据，由浅入深提升分析能力。",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "【案例式教学】",
        "",
        "教师：通过不同案例展示实验动画的多种分类方式，引导学生比较同一作品的多维归类可能性。",
        "",
        "",
        "",
        "",
        "",
        "",
        "【案例式教学】",
        "",
        "教师：引导学生理解材料不仅是媒介，更会影响动画的节奏、肌理和情感表达，完成对“材料语言”的初步认知。",
        "",
        "",
        "【交互式教学】",
        "通过课堂测试检验学习效果",
        "",
        "",
        "",
        "【10分钟】内容总结",
        "教师：引导学生归纳总结；",
        "学生：学到了什么？还存在哪些疑问？引导学生善于总结，发现不足，自主探究。",
        "",
        "",
        "",
        "【10分钟】作业布置",
        "教师：布置“材料语言测试”实践任务，说明作业目标、规格与提交方式，帮助学生明确课后实践方向。",
        "",
        "",
        "",
        "教师：通过指导和督促学生完成课后实践任务，鼓励学生在试错中理解实验动画的探索价值。",
        "",
        "学生：围绕材料观察、运动测试和结果整理开展实践，逐步建立从观看到实验的创作意识。",
        " ",
        "",
        "",
        "",
        "",
        "【课程思政设计】",
        "",
        "",
        "",
        "通过课程内容与作业任务的衔接，引导学生树立创新意识、工匠精神、文化自信和动画创作者的责任感。",
    ]
    fill_existing_paragraphs(main_table.cell(4, 4), design_lines, default_font="宋体", default_size=10.5)
    design_cell = main_table.cell(4, 4)
    design_heading_indexes = {1, 14, 19, 31, 53, 75, 85, 91, 96, 104}
    for idx, paragraph in enumerate(design_cell.paragraphs):
        text = paragraph.text
        if idx in design_heading_indexes:
            for run in paragraph.runs:
                apply_run_style(run, "楷体", 10.5, True)
        elif text.startswith("教师："):
            clear_and_set_segments(
                paragraph,
                [
                    ("教师：", "宋体", 10.5, True),
                    (text[3:], "宋体", 10.5, False),
                ],
            )
        elif text.startswith("学生："):
            clear_and_set_segments(
                paragraph,
                [
                    ("学生：", "宋体", 10.5, True),
                    (text[3:], "宋体", 10.5, False),
                ],
            )
        elif text.startswith("【") and text.endswith("】"):
            for run in paragraph.runs:
                apply_run_style(run, "楷体", 10.5, True)
        elif text.startswith("【") and "分钟" in text:
            for run in paragraph.runs:
                apply_run_style(run, "宋体", 10.5, True)
        else:
            for run in paragraph.runs:
                apply_run_style(run, "宋体", 10.5, False)

    # Keep the original note area and overall page geometry unchanged.
    fill_existing_paragraphs(main_table.cell(6, 0), [""] * len(main_table.cell(6, 0).paragraphs), default_font="宋体", default_size=10.5)

    # Restore bold labels and default fonts in unchanged header cells of the main table.
    fixed_bold_cells = [(0, 0), (0, 2), (0, 4), (1, 0), (3, 0), (3, 4), (5, 0)]
    for row, col in fixed_bold_cells:
        for paragraph in main_table.cell(row, col).paragraphs:
            for run in paragraph.runs:
                apply_run_style(run, "宋体", 10.5, True)

    doc.save(dst)
    print(dst)


if __name__ == "__main__":
    main()
