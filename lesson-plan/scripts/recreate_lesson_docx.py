from __future__ import annotations

import shutil
import sys

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


def set_run_font(run, font_name: str = "宋体", font_size: int = 12) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(font_size)


def set_cell_text(cell, text: str, font_name: str = "宋体", font_size: int = 12) -> None:
    cell.text = ""
    lines = text.split("\n")
    for index, line in enumerate(lines):
        if index == 0:
            paragraph = cell.paragraphs[0]
        else:
            paragraph = cell.add_paragraph()
        run = paragraph.add_run(line)
        set_run_font(run, font_name, font_size)


def normalize_table_fonts(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run)


def main() -> None:
    if len(sys.argv) != 3:
        raise SystemExit("Usage: recreate_lesson_docx.py <template.docx> <output.docx>")

    src, dst = sys.argv[1], sys.argv[2]
    shutil.copyfile(src, dst)
    doc = Document(dst)

    info = doc.tables[1]
    info.cell(0, 1).text = "中国文化实验动画创作"
    info.cell(1, 1).text = "动画艺术学院"
    info.cell(2, 1).text = "2025/2026第二学期"
    info.cell(3, 1).text = "动画"
    info.cell(4, 1).text = ""
    info.cell(5, 1).text = ""
    normalize_table_fonts(info)

    main_table = doc.tables[3]
    main_table.cell(0, 1).text = "1"
    main_table.cell(0, 3).text = "第一章 实验动画概述\n第一课次：实验动画的定义、特征与分类导入"
    main_table.cell(0, 5).text = "4"
    normalize_table_fonts(main_table)

    objective_text = """一、教学目标
知识目标：
1.能够准确说出实验动画的基本定义；
2.能够概括实验动画区别于主流叙事动画的核心特征；
3.能够说出实验动画的三种基本分类维度：按主题、按风格、按材料；
4.能够结合课堂案例，对实验动画作品进行初步分类判断。
能力目标：
1.具备从形式、材料和表达方式切入分析实验动画作品的能力；
2.具备使用基础术语描述实验动画艺术特征的能力；
3.具备围绕“材料语言测试”实践任务进行初步构思的能力。
素养目标：
1.树立勇于创新、敢于突破惯性表达的创作意识；
2.增强对中国动画学派实验传统的认同感与文化自信；
3.培养严谨观察、主动试错和知行合一的学习习惯。
二、重点难点：
重点：实验动画的定义、特征及三种分类维度的理解。
难点：实验动画分类维度之间的交叉关系，以及材料语言如何参与动画表达。
三、学情分析
学生对“动画”已有直观认识，但对“实验动画”的理解仍偏感性，容易只关注题材和画风，忽略形式、材料与表达方式之间的关系。课堂需通过案例对比、分类辨析和材料测试导入，帮助学生建立基础观察框架。"""
    set_cell_text(main_table.cell(2, 0), objective_text)

    content_text = """【课前自主学习】
明确目标，发布任务
在平台或学习群发布课前任务，布置学生通过线上资源进行课前预习。
1.思考什么样的作品可以被称为实验动画；
2.观察实验动画与主流动画的差异；
3.查找1-2个风格或材料较特殊的动画案例。
线上预习，总结归纳
通过线上资源，学生初步形成对实验动画的基础认知，并带着问题进入课堂。
【课中教学内容】
一、课程导入与问题激活
1.说明本课程在本学期中的定位、学习目标和最终成果导向；
2.通过案例对比，引导学生观察主流动画与实验动画的差异；
3.围绕“什么是实验动画”展开课堂讨论。
二、实验动画的定义与特征
1.讲解实验动画的基本定义；
2.梳理实验动画的核心特征：非单一叙事、材料探索、形式创新、个体表达；
3.从叙事目标、视觉逻辑、技术路径和观看体验等方面进行对比分析。
三、实验动画的分类框架
1.讲解按主题、按风格、按材料三种基本分类方式；
2.结合案例说明分类的多维交叉关系；
3.重点引出按材料分类中的二维逐帧与数字合成、综合材料应用、剪纸动画、实物动画、偶动画等内容。
四、中国动画实验传统与课程衔接
1.简要介绍中国动画学派的实验探索传统；
2.说明后续课程将进入中国文化实验动画的探索与创作；
3.本课不展开系统中国文化史知识，而是先建立实验动画方法意识。
五、课堂测试与总结
1.通过课堂问答或选择题检测学习效果；
2.组织学生对案例进行口头分类说明；
3.总结本课知识点。
六、作业布置
内容：完成“材料语言测试”。
要求：利用纸张、布料、树叶等综合材料，或数字笔刷，制作一段约5秒的“材质运动测试”；可附简短文字说明，说明所用材料、表现意图和遇到的问题。"""
    set_cell_text(main_table.cell(4, 0), content_text)

    design_text = """【启发式教学】
教师：课前发布导学任务单，以问题驱动方式引导学生预习“实验动画”的相关内容，激发学生学习兴趣。
学生：完成资料查找、案例收集与预习思考，梳理自己对实验动画的初步理解。
教师：在线答疑，了解学生预习情况，为课堂讲授做准备。
【课中180分钟】
【交互式教学】
【20分钟】课程导入与问题激活
教师：通过案例对比和问题引导，激活学生已有经验。
学生：围绕“什么是实验动画”进行分享与讨论。
【讲解法】
【45分钟】实验动画的定义与特征
教师：系统讲解实验动画的定义与核心特征，并与主流动画进行比较。
学生：记录关键词，结合案例判断不同特征的体现方式。
【案例式教学】
【55分钟】实验动画的分类框架
教师：围绕按主题、按风格、按材料三类分类进行讲解，并组织案例辨析。
学生：分组完成案例分类练习，说明分类依据。
【课程思政设计】
教师：在讲解实验精神时，引导学生认识创新不是脱离传统，而是在理解传统基础上的创造性转化。
教师：结合中国动画学派的探索传统，增强学生文化自信和民族动画认同。
【案例式教学】
【20分钟】中国动画实验传统与课程衔接
教师：简要介绍水墨、剪纸、折纸等方向的实验探索，说明后续课程发展脉络。
【交互式教学】
【20分钟】课堂练习与即时检测
教师：通过课堂测试与口头表达检验学生是否掌握定义、特征与分类框架。
学生：完成测试并进行案例分类说明。
【10分钟】内容总结
教师：引导学生归纳本课知识点，梳理“定义-特征-分类-任务”的结构关系。
学生：总结学习收获与疑问。
【10分钟】作业布置
教师：布置“材料语言测试”任务，说明目标、规格与提交方式，强调实验意识与材料观察的重要性。
学生：根据任务要求进行课后准备与实践。
【课程思政设计】
通过课程内容与作业任务的衔接，引导学生树立创新意识、工匠精神、文化自信和动画创作者的责任感。"""
    set_cell_text(main_table.cell(4, 4), design_text)

    set_cell_text(main_table.cell(6, 0), "\n\n\n\n\n\n\n\n\n\n\n")
    normalize_table_fonts(main_table)

    doc.save(dst)
    print(dst)


if __name__ == "__main__":
    main()
